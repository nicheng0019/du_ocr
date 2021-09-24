import cv2
from dulines import *
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
import docx

PAPER_HEIGHT = 22.1
PAPER_WIDTH = 15.2


def get_top_line(lines, imgh):
    ymin = imgh
    topidx = -1
    for li, line in enumerate(lines):
        if not line.valid:
            continue

        if line.vertical:
            continue

        if line.center[1] < ymin:
            ymin = line.center[1]
            topidx = li

    return topidx


def points_dist(p1, p2):
    return np.linalg.norm(p1 - p2)


def get_rect(lines, topidx):
    topline = lines[topidx]
    topp1 = topline.p1
    topp2 = topline.p2

    rects = []

    linenum = len(lines)
    for leftidx in range(linenum):
        leftline = lines[leftidx]
        if not leftline.valid:
            continue

        if not leftline.vertical:
            continue

        leftp1 = leftline.p1
        leftp2 = leftline.p2

        if points_dist(leftp1, topp1) > 20:
            continue

        for bottomidx in range(linenum):
            bottomline = lines[bottomidx]
            if not bottomline.valid:
                continue

            if bottomline.vertical:
                continue

            if bottomidx == topidx:
                continue

            bottomp1 = bottomline.p1
            bottomp2 = bottomline.p2

            if points_dist(leftp2, bottomp1) > 20:
                continue

            for rightidx in range(linenum):
                rightline = lines[rightidx]
                if not rightline.valid:
                    continue

                if not rightline.vertical:
                    continue

                if leftidx == rightidx:
                    continue

                rightp1 = rightline.p1
                rightp2 = rightline.p2

                if points_dist(rightp2, bottomp2) > 20:
                    continue

                if points_dist(rightp1, topp2) > 20:
                    continue

                rects.append([topidx, leftidx, bottomidx, rightidx])

    return rects


def choose_rect(lines, rects):
    maxarea = 0
    maxidx = -1
    for ri, rect in enumerate(rects):
        top, left, bottom, right = rect
        topline, leftline, bottomline, rightline = lines[top], lines[left], lines[bottom], lines[right]
        area = (bottomline.center[1] - topline.center[1]) * (rightline.center[0] - leftline.center[0])

        if area > maxarea:
            maxarea = area
            maxidx = ri

    return rects[maxidx]


def get_all_lines_in_rect(lines, rect):
    top, left, bottom, right = rect
    topline, leftline, bottomline, rightline = lines[top], lines[left], lines[bottom], lines[right]

    bound = [topline.center[1], leftline.center[0], bottomline.center[1], rightline.center[0]]
    sub_lines = []
    sub_lines.extend(rect)
    for li, line in enumerate(lines):
        if not line.valid:
            continue

        if line.in_bound(bound):
            sub_lines.append(li)

    return sub_lines


def get_sub_table(lines, sub_lines):
    THRESH = 10
    topline, leftline, bottomline, rightline = lines[sub_lines[0]], \
                                               lines[sub_lines[1]], lines[sub_lines[2]], lines[sub_lines[3]]

    horizontal_lines_id = []
    vertical_lines_id = []
    for li in sub_lines[4:]:
        if lines[li].vertical:
            vertical_lines_id.append(li)
        else:
            horizontal_lines_id.append(li)

    row_num = column_num = 1

    prevrow = topline.center[1]
    row_ys = defaultdict(list)
    row_ys[row_num - 1].append(topline.center[1])

    horizontal_id_dict = defaultdict(list)
    for li in horizontal_lines_id:
        line = lines[li]

        if line.center[1] - prevrow > THRESH:
            row_num += 1

        lines[li].rowstart = lines[li].rowend = row_num - 1
        horizontal_id_dict[row_num - 1].append(li)

        prevrow = line.center[1]
        row_ys[row_num - 1].append(line.center[1])

    prevcol = leftline.center[0]
    col_xs = defaultdict(list)
    col_xs[column_num - 1].append(leftline.center[0])

    vertical_id_dict = defaultdict(list)
    for li in vertical_lines_id:
        line = lines[li]
        if line.center[0] - prevcol > THRESH:
            column_num += 1

        lines[li].colstart = lines[li].colend = column_num - 1
        vertical_id_dict[column_num - 1].append(li)

        prevcol = line.center[0]
        col_xs[column_num - 1].append(line.center[0])

    row_num += 1
    column_num += 1
    row_ys[row_num - 1].append(bottomline.center[1])
    col_xs[column_num - 1].append(rightline.center[0])
    print("row_num: ", row_num, "column_num: ", column_num)

    row_ys_arr = np.zeros(len(row_ys.keys()))
    for key in row_ys.keys():
        row_ys_arr[key] = sum(row_ys[key]) / len(row_ys[key])

    col_xs_arr = np.zeros(len(col_xs.keys()))
    for key in col_xs.keys():
        col_xs_arr[key] = sum(col_xs[key]) / len(col_xs[key])

    for li in horizontal_lines_id:
        lines[li].colstart = np.argmin(np.abs(lines[li].p1[0] - col_xs_arr))
        lines[li].colend = np.argmin(np.abs(lines[li].p2[0] - col_xs_arr))

    for li in vertical_lines_id:
        lines[li].rowstart = np.argmin(np.abs(lines[li].p1[1] - row_ys_arr))
        lines[li].rowend = np.argmin(np.abs(lines[li].p2[1] - row_ys_arr))

    lines[sub_lines[0]].rowstart = lines[sub_lines[0]].rowend = 0
    lines[sub_lines[2]].rowstart = lines[sub_lines[2]].rowend = row_num - 1

    lines[sub_lines[0]].colstart = lines[sub_lines[2]].colstart = 0
    lines[sub_lines[0]].colend = lines[sub_lines[2]].colend = column_num - 1

    lines[sub_lines[1]].colstart = lines[sub_lines[1]].colend = 0
    lines[sub_lines[3]].colstart = lines[sub_lines[3]].colend = column_num - 1

    lines[sub_lines[1]].rowstart = lines[sub_lines[3]].rowstart = 0
    lines[sub_lines[1]].rowend = lines[sub_lines[3]].rowend = row_num - 1

    cell_horizontal = np.repeat(np.arange(column_num - 1)[np.newaxis, :], row_num - 1, axis=0)
    cell_vertical = np.repeat(np.arange(row_num - 1)[np.newaxis, :], column_num - 1, axis=0)

    for i in range(1, row_num - 1):
        for j in range(1, column_num - 1):
            vertical_ids = vertical_id_dict[j]
            cols_segs = [[lines[li].rowstart,
                          lines[li].rowend] for li in vertical_ids]

            if len(cols_segs) == 0:
                continue

            for n in range(0, cols_segs[0][0]):
                cell_horizontal[n, j - 1] = j

            for s in range(1, len(cols_segs)):
                for n in range(cols_segs[s - 1][1], cols_segs[s][0]):
                    cell_horizontal[n, j - 1] = j

            for n in range(cols_segs[-1][1], row_num - 1):
                cell_horizontal[n, j - 1] = j

    for i in range(1, column_num - 1):
        for j in range(1, row_num - 1):
            horizontal_ids = horizontal_id_dict[j]
            rows_segs = [[lines[li].colstart,
                          lines[li].colend] for li in horizontal_ids]

            if len(rows_segs) == 0:
                continue

            for n in range(0, rows_segs[0][0]):
                cell_vertical[n, j - 1] = j

            for s in range(1, len(rows_segs)):
                for n in range(rows_segs[s - 1][1], rows_segs[s][0]):
                    cell_vertical[n, j - 1] = j

            for n in range(rows_segs[-1][1], column_num - 1):
                cell_vertical[n, j - 1] = j

    return row_num, column_num, cell_horizontal, cell_vertical, row_ys_arr, col_xs_arr


def docx_gen(lines, img, boxes, txts):
    imgh, imgw = img.shape[:2]

    lines_arr = []
    for line in lines:
        lines_arr.append([line.p1, line.p2])

    avgangle = 0
    count = 0
    for line in lines:
        if not line.vertical:
            avgangle += line.angle
            count += 1

    if count == 0:
        for line in lines:
            if line.vertical:
                if line.angle > 0:
                    avgangle += line.angle - 90
                else:
                    avgangle += line.angle + 90
                count += 1

    assert count != 0
    avgangle = avgangle / count
    print("avgangle: ", avgangle)

    center = (imgw / 2, imgh / 2)
    M = cv2.getRotationMatrix2D(center, avgangle, 1)
    lines_arr = np.array(lines_arr)
    lines_arr = lines_arr.dot(M[:, :2].T) + M[:, 2]

    boxes_arr = np.array(boxes, dtype=np.float32)
    boxes_arr = boxes_arr.dot(M[:, :2].T) + M[:, 2]

    img = cv2.warpAffine(img, M, (img.shape[1], img.shape[0]))

    lines = []
    for line in lines_arr:
        lines.append(Line(line))

    lines.sort()

    document = Document()
    document.styles["Normal"].font.name = u"宋体"
    document.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
    document.styles["Normal"].font.size = Pt(10)

    while True:
        topidx = get_top_line(lines, imgh)
        if -1 == topidx:
            break

        rects = get_rect(lines, topidx)
        for rect in rects:
            lines[rect[0]].valid = False
            lines[rect[1]].valid = False
            lines[rect[2]].valid = False
            lines[rect[3]].valid = False

        rect = choose_rect(lines, rects)

        sub_lines = get_all_lines_in_rect(lines, rect)
        for li in sub_lines[4:]:
            lines[li].valid = False
            cv2.line(img, tuple(lines[li].p1.astype(np.int32)), tuple(lines[li].p2.astype(np.int32)),
                     (255, 0, 0), 10)

        # cell_horizontal: table_row_num*table_column_num, cell_vertical: table_column_num*table_row_num
        row_num, column_num, cell_horizontal, cell_vertical, row_ys_arr, col_xs_arr = get_sub_table(lines, sub_lines)

        table_row_num = row_num - 1
        table_column_num = column_num - 1
        table = document.add_table(rows=table_row_num, cols=table_column_num, style='Table Grid')
        table.style.paragraph_format.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

        resolutionH = PAPER_HEIGHT / (row_ys_arr[-1] - row_ys_arr[0])
        resolutionW = PAPER_WIDTH / (col_xs_arr[-1] - col_xs_arr[0])

        cells_pos = np.zeros((table_row_num, table_column_num, 4), dtype=np.float32)

        for row in range(table_row_num):
            table.rows[row].height = Cm((row_ys_arr[row + 1] - row_ys_arr[row]) * resolutionH)
            for col in range(table_column_num):
                table.cell(row, col).width = Cm((col_xs_arr[col + 1] - col_xs_arr[col]) * resolutionW)

                cells_pos[row, col] = np.array([col_xs_arr[col],
                                                col_xs_arr[col + 1], row_ys_arr[row], row_ys_arr[row + 1]])

        for row in range(table_row_num):
            lastcol = 0
            for col in range(table_column_num):
                if cell_horizontal[row, col] != col:
                    continue

                if col == lastcol == cell_horizontal[row, col]:
                    lastcol = col + 1
                    continue

                table.cell(row, lastcol).merge(table.cell(row, col))
                cells_pos[row, lastcol, 1] = cells_pos[row, col, 1]
                lastcol = col + 1

        for col in range(table_column_num):
            lastrow = 0
            for row in range(table_row_num):
                if cell_vertical[col, row] != row:
                    continue

                if row == lastrow == cell_vertical[col, row]:
                    lastrow = row + 1
                    continue

                table.cell(lastrow, col).merge(table.cell(row, col))
                cells_pos[lastrow, col, 3] = cells_pos[row, col, 3]
                lastrow = row + 1

        cell_text_dict = defaultdict(list)
        for bi, box in enumerate(boxes_arr):
            boxleft = np.min(box[:, 0])
            boxright = np.max(box[:, 0])
            boxtop = np.min(box[:, 1])
            boxbottom = np.max(box[:, 1])

            bfind = False
            for row in range(table_row_num):
                for col in range(table_column_num):
                    if boxleft >= cells_pos[row, col, 0] and boxright <= cells_pos[row, col, 1]\
                            and boxtop >= cells_pos[row, col, 2] and boxbottom <= cells_pos[row, col, 3]:
                        bfind = True

                        cell_text_dict[(row, col)].append(txts[bi])
                        break

                if bfind:
                    break

        for row in range(table_row_num):
            for col in range(table_column_num):
                if len(cell_text_dict[(row, col)]) == 0:
                    continue

                text = " ".join(cell_text_dict[(row, col)])
                table.cell(row, col).text = text
                table.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    document.save('d:/demo.docx')
    quit()


if __name__ == "__main__":
    linefn, imgfn = r"D:\Dataset\ocr\0002_result.txt", r"D:\Dataset\ocr\0002.jpg"

    # lines = read_line_file(linefn)
    # img = cv2.imread(imgfn)
    # docx_gen(lines, img)

    document = Document()
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    table = document.add_table(rows=4, cols=5, style='Table Grid')
    document.styles["Normal"].font.name = u"宋体"
    document.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
    document.styles["Normal"].font.size = Pt(10.5)

    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(1, 0).merge(table.cell(1, 1))
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 1).merge(table.cell(1, 1))
    hdr_cells0 = table.rows[0].cells
    table.rows[0].height = Cm(3)
    p = hdr_cells0[2].add_paragraph('一二三四五六七八九十十一十二')
    hdr_cells0[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    hdr_cells0[3].add_paragraph('3\n')
    document.save('d:/demo.docx')
